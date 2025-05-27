# portus_sa_module/sa_report.py

from pathlib import Path
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from portus_sa_module.sa_engine import score_all_reviews


def generate_sa_report(clean_file: Path, provider: str, output_folder: Path) -> None:
    if not clean_file.exists():
        raise FileNotFoundError(clean_file)

    reply_col = {"openai": "Reply OAI", "google": "Reply Go", "xai": "Reply XAI"}[provider]

    wb = load_workbook(clean_file)
    ws = wb.active
    headers = [c.value for c in ws[1]]

    def resolve_idx(*names):
        for name in names:
            if name in headers:
                return headers.index(name)
        return None

    pub_i  = resolve_idx("publishedDate", "reviewDate", "publishedAtDate")
    rate_i = resolve_idx("rating", "reviewScoreWithDescription/label")
    tit_i  = resolve_idx("title", "reviewTitle")
    user_i = resolve_idx("user/name", "userName", "name")
    txt_i  = resolve_idx("text")
    sa_i   = resolve_idx("sa_score")
    rep_i  = resolve_idx(reply_col)

    entries = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        title = row[tit_i] if tit_i is not None else ""
        text  = row[txt_i] if txt_i is not None else ""
        if str(title).strip() or str(text).strip():
            entries.append({"title": title, "text": text})

    overall, summary = score_all_reviews(entries)

    doc = Document()
    doc.styles['Normal'].font.size = Pt(11)
    doc.add_heading("Sentiment Analysis and Summary", level=1)
    doc.add_paragraph()

    p = doc.add_paragraph()
    run_label = p.add_run("Global Score: ")
    run_label.bold = True
    p.add_run(str(overall))

    doc.add_paragraph(summary)
    doc.add_paragraph()

    doc.add_heading("Row Details", level=1)
    doc.add_paragraph()
    for row_num, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=1):
        title = row[tit_i] if tit_i is not None else ""
        text  = row[txt_i] if txt_i is not None else ""
        if not (str(title).strip() or str(text).strip()):
            continue

        p_row = doc.add_paragraph()
        run_rnum = p_row.add_run(f"Row #: ")
        run_rnum.bold = True
        p_row.add_run(str(row_num))

        p_sa = doc.add_paragraph()
        run_sa = p_sa.add_run("SA Score: ")
        run_sa.bold = True
        p_sa.add_run(str(row[sa_i] if sa_i is not None else 'N/A'))

        p_pub = doc.add_paragraph()
        run_pub = p_pub.add_run("Published: ")
        run_pub.bold = True
        p_pub.add_run(str(row[pub_i] if pub_i is not None else 'N/A'))

        p_rate = doc.add_paragraph()
        run_rate = p_rate.add_run("Rating: ")
        run_rate.bold = True
        p_rate.add_run(str(row[rate_i] if rate_i is not None else 'N/A'))

        p_title = doc.add_paragraph()
        run_title = p_title.add_run("Title: ")
        run_title.bold = True
        p_title.add_run(str(title))

        p_user = doc.add_paragraph()
        run_user = p_user.add_run("User: ")
        run_user.bold = True
        p_user.add_run(str(row[user_i] if user_i is not None else 'N/A'))

        p_txt = doc.add_paragraph()
        run_txt = p_txt.add_run("Text: ")
        run_txt.bold = True
        p_txt.add_run(f"\n{str(text)}")

        p_reply = doc.add_paragraph()
        run_reply = p_reply.add_run(f"{reply_col}: ")
        run_reply.bold = True
        p_reply.add_run(f"\n{str(row[rep_i] if rep_i is not None else 'N/A')}")

        doc.add_paragraph()

    name = clean_file.stem.replace("Clean", "SA") + ".docx"
    out  = output_folder / name
    doc.save(out)
